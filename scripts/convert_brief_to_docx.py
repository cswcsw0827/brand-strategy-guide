#!/usr/bin/env python3
"""Convert a brief markdown file to a formatted Word document (.docx).

Usage:
    python convert_brief_to_docx.py input.md [output.docx]

If no output path is given, the output file is created next to the input file
with the same name and a .docx extension.
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    tcPr.append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, font_size=None, color=None, alignment=None, space_after=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    run = p.add_run(text)
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_inline_formatted_paragraph(doc, segments, style='Normal', alignment=None):
    """Add a paragraph with inline formatting. segments is a list of (text, bold, italic) tuples."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def parse_inline_formatting(text):
    """Parse inline markdown formatting: **bold**, *italic*, `code`."""
    segments = []
    # Pattern for bold, italic, and code
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                segments.append((plain, False, False))

        if match.group(2) is not None:  # bold
            segments.append((match.group(2), True, False))
        elif match.group(3) is not None:  # italic
            segments.append((match.group(3), False, True))
        elif match.group(4) is not None:  # code
            segments.append((match.group(4), False, False))

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        segments.append((text[last_end:], False, False))

    return segments if segments else [(text, False, False)]

def add_rich_paragraph(doc, text, style='Normal', space_after=None):
    """Add a paragraph with inline formatting support."""
    p = doc.add_paragraph(style=style)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    segments = parse_inline_formatting(text)
    for text_seg, bold, italic in segments:
        run = p.add_run(text_seg)
        run.bold = bold
        run.italic = italic
    return p

def is_table_separator(line):
    """Check if a line is a markdown table separator like |---|---|"""
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))

def parse_table_row(line):
    """Parse a markdown table row into cells."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]

def is_table_start(lines, i):
    """Check if current line starts a markdown table."""
    line = lines[i].strip()
    if not line.startswith('|'):
        return False
    if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
        return True
    return False

def add_table_from_markdown(doc, table_lines):
    """Add a Word table from markdown table lines."""
    # Filter out separator lines
    data_lines = [l for l in table_lines if not is_table_separator(l)]
    if not data_lines:
        return

    # Parse all rows
    rows = [parse_table_row(line) for line in data_lines]
    # Normalize column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    # Create Word table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            # Clear default paragraph
            cell.paragraphs[0].clear()

            # Parse inline formatting in cell
            segments = parse_inline_formatting(cell_text)
            for text_seg, bold, italic in segments:
                run = cell.paragraphs[0].add_run(text_seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(8)

            # Header row shading
            if i == 0:
                set_cell_shading(cell, '2B579A')
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True

    doc.add_paragraph()  # spacing after table

def convert_markdown_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    in_blockquote = False
    blockquote_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(60, 60, 60)
                # Add grey background via shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F5F5')
                shd.set(qn('w:val'), 'clear')
                pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if is_table_start(lines, i):
            in_table = True
            table_lines = []
            # Collect table lines
            while i < len(lines):
                l = lines[i].strip()
                if l.startswith('|') and '|' in l[1:]:
                    table_lines.append(lines[i])
                    i += 1
                else:
                    break
            # Render table
            add_table_from_markdown(doc, table_lines)
            in_table = False
            continue

        # Blockquote handling
        if stripped.startswith('>'):
            in_blockquote = True
            blockquote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                blockquote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            # Render blockquote
            for bl in blockquote_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(bl)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.italic = True
            in_blockquote = False
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headers
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            add_rich_paragraph(doc, text, style='Heading 1')
            i += 1
            continue
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            add_rich_paragraph(doc, text, style='Heading 2')
            i += 1
            continue
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            text = stripped[4:]
            add_rich_paragraph(doc, text, style='Heading 3')
            i += 1
            continue
        if stripped.startswith('#### '):
            text = stripped[5:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # List items
        if re.match(r'^[\-\*]\s+', stripped):
            text = re.sub(r'^[\-\*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            segments = parse_inline_formatting(text)
            p.clear()
            for text_seg, bold, italic in segments:
                run = p.add_run(text_seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10)
            i += 1
            continue

        # Checkbox items
        if stripped.startswith('- [x]') or stripped.startswith('- [ ]'):
            checked = stripped.startswith('- [x]')
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run('☑ ' if checked else '☐ ')
            run.font.size = Pt(10)
            segments = parse_inline_formatting(text)
            for text_seg, bold, italic in segments:
                run = p.add_run(text_seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10)
            i += 1
            continue

        # Numbered list items
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            segments = parse_inline_formatting(text)
            p.clear()
            for text_seg, bold, italic in segments:
                run = p.add_run(text_seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10)
            i += 1
            continue

        # Bold header lines (like **角色一：雾人**)
        if stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 120:
            text = stripped[2:-2]
            add_rich_paragraph(doc, f'**{text}**', space_after=2)
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, stripped)
        i += 1

    # Save
    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    return docx_path

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python convert_brief_to_docx.py input.md [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        base = os.path.splitext(input_file)[0]
        output_file = base + '.docx'

    convert_markdown_to_docx(input_file, output_file)
